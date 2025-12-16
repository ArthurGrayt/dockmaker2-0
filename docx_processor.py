import os
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
import shutil

class DocxProcessor:
    def __init__(self, upload_folder, output_folder):
        self.upload_folder = upload_folder
        self.output_folder = output_folder
        os.makedirs(output_folder, exist_ok=True)

    def process_document(self, content_path, model_path, output_filename):
        # Strategy:
        # 1. Load Model Document (to keep headers/footers/styles).
        # 2. Clear its body content.
        # 3. Load Content Document.
        # 4. Copy text/formatting from Content to Model.
        
        # Load Model
        doc = Document(model_path)
        
        # Clear existing body content
        self._clear_body(doc)
        
        # Load Source Content
        source_doc = Document(content_path)
        
        # Extract and Append Content
        self._copy_content(source_doc, doc)
        
        # Ensure Section Properties (Headers/Footers default to Model's)
        # We need to ensure that if the source had multiple sections, 
        # we might want to respect them or merge them.
        # Requirement: "TODAS as páginas utilizam o cabeçalho... do modelo".
        # So we effectively enforce the Model's section props on the whole doc.
        # But if the source text is long, it will naturally flow.
        
        output_path = os.path.join(self.output_folder, output_filename)
        doc.save(output_path)
        return output_path

    def _clear_body(self, doc):
        """Removes all paragraphs and tables from the document body."""
        body = doc._body
        body.clear_content()

    def _copy_content(self, source, target):
        """Copies paragraphs and run properties from source to target."""
        for para in source.paragraphs:
            if not para.text.strip() and len(para.runs) == 0:
                continue # Skip completely empty paragraphs? No, might be spacing.
                
            new_p = target.add_paragraph()
            
            # Copy Paragraph Alignment
            if para.alignment:
                new_p.alignment = para.alignment
            
            # Copy Paragraph Style? 
            # Ideally keep 'Normal' or map it, but let's stick to basic formatting as requested.
            
            for run in para.runs:
                new_r = new_p.add_run(run.text)
                new_r.bold = run.bold
                new_r.italic = run.italic
                new_r.underline = run.underline
                # Copy Font Size/Name if present?
                # "Negrito, itálico e alinhamento básico" is the requirement.
                # If we copy everything, we risk clashing with Model styles.
                # But let's copy explicit font sizes if they exist.
                if run.font.size:
                    new_r.font.size = run.font.size
                if run.font.name:
                    new_r.font.name = run.font.name
                
            # Handle Line Breaks? 
            # python-docx paragraphs handle basic wrapping.
            
    def get_preview_info(self, model_path):
        """Returns info for preview."""
        return os.path.basename(model_path)
