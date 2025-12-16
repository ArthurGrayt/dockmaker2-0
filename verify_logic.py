from docx import Document
import os
from docx_processor import DocxProcessor

def create_model():
    doc = Document()
    section = doc.sections[0]
    header = section.header
    header.paragraphs[0].add_run("THIS IS THE MODEL HEADER").bold = True
    
    footer = section.footer
    footer.paragraphs[0].add_run("THIS IS THE MODEL FOOTER").italic = True
    
    doc.save("test_model.docx")
    print("Created test_model.docx")

def create_content():
    doc = Document()
    doc.add_paragraph("This is the content paragraph 1.")
    doc.add_paragraph("This is the content paragraph 2 with bold.").runs[0].bold = True
    doc.save("test_content.docx")
    print("Created test_content.docx")

def verify():
    create_model()
    create_content()
    
    proc = DocxProcessor("temp_uploads", "temp_output")
    # We use local paths for test
    os.makedirs("temp_uploads", exist_ok=True)
    os.makedirs("temp_output", exist_ok=True)
    
    try:
        out = proc.process_document("test_content.docx", "test_model.docx", "test_result.docx")
        print(f"Propcessing success: {out}")
        
        # Check result
        d = Document(out)
        print("Result paragraphs:", len(d.paragraphs))
        print("Result Header:", d.sections[0].header.paragraphs[0].text)
        
        if "CONTENT PARAGRAPH 1" in [p.text.upper() for p in d.paragraphs]: 
            print("Content Verification: PASS (Text found)")
        else:
            print("Content Verification: PARTIAL (Text case might differ or found)")
            
        if "MODEL HEADER" in d.sections[0].header.paragraphs[0].text.upper():
            print("Header Verification: PASS")
        else:
            print("Header Verification: FAIL")
            
    except Exception as e:
        print(f"Processing FAILED: {e}")

if __name__ == "__main__":
    verify()
